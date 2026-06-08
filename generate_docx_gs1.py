import json
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

# Load GS1 questions
with open("gs1_questions.json", "r", encoding="utf-8") as f:
    questions = json.load(f)

# Reference taxonomy mappings (same as script 2)
from generate_final_gs1_v2 import get_broad_section_group, generate_short_desc, determine_tag

# Group questions: Subject -> Broad Section Group -> Topic -> Microtheme
grouped = {}
for q in questions:
    subject = q["subject"]
    topic = q["topic"]
    microtheme = q["microtheme"]
    
    broad_section = get_broad_section_group(topic)
    
    if subject not in grouped:
        grouped[subject] = {}
    if broad_section not in grouped[subject]:
        grouped[subject][broad_section] = {}
    if topic not in grouped[subject][broad_section]:
        grouped[subject][broad_section][topic] = {}
    if microtheme not in grouped[subject][broad_section][topic]:
        grouped[subject][broad_section][topic][microtheme] = []
        
    pyq = q["question"]
    marks = q["marks"]
    year = q["year"]
    
    short_desc = generate_short_desc(pyq, microtheme)
    tag = determine_tag(pyq, marks)
    
    grouped[subject][broad_section][topic][microtheme].append({
        "question": pyq,
        "year": year,
        "marks": marks,
        "short_desc": short_desc,
        "tag": tag
    })

# Initialize Word Document
doc = Document()

# Page Setup (Landscape for tables)
section = doc.sections[0]
section.orientation = 1  # 1 is Landscape
# Swap width and height for landscape
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height
# Set 0.5-inch margins for more space
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Styling Helpers
COLOR_PRIMARY = RGBColor(27, 54, 93)   # Deep Navy (#1B365D)
COLOR_TEXT = RGBColor(51, 51, 51)      # Charcoal
COLOR_MUTED = RGBColor(119, 119, 119)  # Medium Grey
HEX_PRIMARY = "1B365D"
HEX_ZEBRA = "F7F9FB"
HEX_BORDER = "D3D3D3"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>')
    tcPr.append(shading)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'  <w:top w:val="single" w:sz="6" w:space="0" w:color="1B365D"/>'
        r'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="1B365D"/>'
        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        r'  <w:left w:val="none"/>'
        r'  <w:right w:val="none"/>'
        r'  <w:insideV w:val="none"/>'
        r'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_table_padding(table):
    tblPr = table._tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        mar = OxmlElement(f'w:{side}')
        # Top/Bottom padding: 8 pt (160 dxa), Left/Right: 10 pt (200 dxa)
        val = '160' if side in ['top', 'bottom'] else '200'
        mar.set(qn('w:w'), val)
        mar.set(qn('w:type'), 'dxa')
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)

# Title Page / Document Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("UPSC GS Mains Paper 1 (GS-I) Syllabus Hierarchy & Tagged PYQs")
run.font.name = "Arial"
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = COLOR_PRIMARY

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = subtitle_p.add_run("Complete Question Bank (2013-2025) structured in 6 layers with directive-based behavioral tags")
run_sub.font.name = "Arial"
run_sub.font.size = Pt(12)
run_sub.font.italic = True
run_sub.font.color.rgb = COLOR_MUTED

# Add a spacer
doc.add_paragraph().paragraph_format.space_after = Pt(20)

# Build Content
for subject in sorted(grouped.keys()):
    # Subject Heading
    h_subj = doc.add_paragraph()
    h_subj.paragraph_format.space_before = Pt(24)
    h_subj.paragraph_format.space_after = Pt(6)
    h_subj.paragraph_format.keep_with_next = True
    run_subj = h_subj.add_run(f"SUBJECT: {subject}")
    run_subj.font.name = "Arial"
    run_subj.font.size = Pt(16)
    run_subj.font.bold = True
    run_subj.font.color.rgb = COLOR_PRIMARY
    
    for bs in sorted(grouped[subject].keys()):
        # Broad Section Group Heading
        h_bs = doc.add_paragraph()
        h_bs.paragraph_format.space_before = Pt(14)
        h_bs.paragraph_format.space_after = Pt(10)
        h_bs.paragraph_format.keep_with_next = True
        run_bs = h_bs.add_run(f"Broad Section Group: {bs}")
        run_bs.font.name = "Arial"
        run_bs.font.size = Pt(13)
        run_bs.font.bold = True
        run_bs.font.color.rgb = RGBColor(90, 90, 90)
        
        # Create Table for this Section Group
        # Columns: 1. Syllabus Point | 2. Microtheme | 3. Question Details | 4. Short Desc | 5. Tag
        table = doc.add_table(rows=1, cols=5)
        table.autofit = False
        set_table_borders(table)
        set_table_padding(table)
        
        # Column Widths (total page width is ~10 inches after margins)
        widths = [Inches(1.8), Inches(1.4), Inches(4.6), Inches(1.4), Inches(0.8)]
        
        # Header Row
        hdr_cells = table.rows[0].cells
        headers = ["Syllabus Point", "Microtheme", "Question Details & Path", "Short Desc (L6)", "Tag"]
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            hdr_cells[idx].width = widths[idx]
            set_cell_background(hdr_cells[idx], HEX_PRIMARY)
            # Text style
            p = hdr_cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.runs[0]
            r.font.name = "Arial"
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        # Enable Header repeating
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:tblHeader'))
        
        # Populate Table Rows
        row_idx = 0
        for topic in sorted(grouped[subject][bs].keys()):
            for microtheme in sorted(grouped[subject][bs][topic].keys()):
                for q_data in grouped[subject][bs][topic][microtheme]:
                    row_idx += 1
                    row = table.add_row()
                    
                    # Prevent row splitting
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(OxmlElement('w:cantSplit'))
                    
                    cells = row.cells
                    # Alternate background coloring
                    bg_color = HEX_ZEBRA if row_idx % 2 == 1 else "FFFFFF"
                    
                    # 1. Syllabus Point
                    cells[0].text = topic
                    # 2. Microtheme
                    cells[1].text = microtheme
                    
                    # 3. Question Details
                    q_p = cells[2].paragraphs[0]
                    q_p.paragraph_format.space_after = Pt(4)
                    q_run = q_p.add_run(q_data["question"])
                    q_run.font.name = "Arial"
                    q_run.font.size = Pt(9.5)
                    q_run.font.color.rgb = COLOR_TEXT
                    
                    # Add Year/Marks as bold line
                    meta_p = cells[2].add_paragraph()
                    meta_p.paragraph_format.space_after = Pt(2)
                    meta_run = meta_p.add_run(f"Year: {q_data['year']} | Marks: {q_data['marks']}")
                    meta_run.font.name = "Arial"
                    meta_run.font.size = Pt(8.5)
                    meta_run.font.bold = True
                    meta_run.font.color.rgb = COLOR_PRIMARY
                    
                    # Add Path line
                    topic_name = topic.split('.', 1)[1].strip() if '.' in topic else topic
                    path = f"GS-I ➔ {subject} ➔ {bs} ➔ {topic_name} ➔ {microtheme} ➔ {q_data['short_desc']}"
                    path_p = cells[2].add_paragraph()
                    path_p.paragraph_format.space_after = Pt(0)
                    path_run = path_p.add_run(f"Path: {path}")
                    path_run.font.name = "Arial"
                    path_run.font.size = Pt(7.5)
                    path_run.font.italic = True
                    path_run.font.color.rgb = COLOR_MUTED
                    
                    # 4. Short Description
                    cells[3].text = q_data["short_desc"]
                    
                    # 5. Behavioral Tag
                    cells[4].text = q_data["tag"]
                    
                    # Set formatting for all cells in the row
                    for c_idx, cell in enumerate(cells):
                        cell.width = widths[c_idx]
                        set_cell_background(cell, bg_color)
                        # Except column 3 which we formatted manually, set fonts for basic text cells
                        if c_idx in [0, 1, 3, 4]:
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(0)
                            if p.runs:
                                r = p.runs[0]
                                r.font.name = "Arial"
                                r.font.size = Pt(9)
                                r.font.color.rgb = COLOR_TEXT
                                if c_idx == 4: # Highlight the tag
                                    r.font.bold = True
                                    if q_data["tag"] == "Short Note":
                                        r.font.color.rgb = RGBColor(180, 80, 0)
                                    elif q_data["tag"] == "Applied":
                                        r.font.color.rgb = RGBColor(0, 120, 80)

# Save Document
out_dir = r"c:\Users\Dr. Yogesh\Desktop\mains\neet and upsc cms\syllabus hierarchy"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "GS1_Syllabus_Hierarchy.docx")
doc.save(out_path)

print(f"Generated DOCX Syllabus Hierarchy at: {out_path}")
