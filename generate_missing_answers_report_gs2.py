import os
import re
import docx
from docx.shared import Inches, RGBColor, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Base Directories
syllabus_dir = r"C:\Users\Dr. Yogesh\Desktop\mains\neet and upsc cms\upsc\syllabus hierarchy\gs2"
solved_dir = r"C:\Users\Dr. Yogesh\Desktop\mains\neet and upsc cms\upsc\solved paper\gs2"

institute_files = {
    "Civilsdaily": os.path.join(solved_dir, "gs2_civilsdaily.md"),
    "Drishti IAS": os.path.join(solved_dir, "gs2_drishti_ias.md"),
    "PWOnlyIAS": os.path.join(solved_dir, "gs2_pwonlyias.md"),
    "Superkalam": os.path.join(solved_dir, "gs2_superkalam.md"),
    "Unacademy": os.path.join(solved_dir, "gs2_unacademy.md")
}

def clean_and_tokenize(text):
    text = text.lower()
    text = re.sub(r'^(?:q\d+\.?|que\.?|question\s*\d+\.?|answer\s*in\s*\d+\s*words|marks?|words?|\d+\s*marks?|\d+\s*words?)\s*', '', text)
    text = re.sub(r'[^a-z0-9\s]', '', text)
    tokens = set(text.split())
    stop_words = {'the', 'and', 'of', 'to', 'in', 'is', 'that', 'it', 'on', 'with', 'as', 'for', 'was', 'were'}
    return tokens - stop_words

def jaccard_similarity(set1, set2):
    if not set1 or not set2:
        return 0.0
    return len(set1.intersection(set2)) / len(set1.union(set2))

def parse_institute_files():
    institute_data = {}
    for inst_name, path in institute_files.items():
        if not os.path.exists(path):
            print(f"Warning: Institute file not found: {path}")
            continue
        
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        
        # Split by `## Question`
        blocks = re.split(r'\r?\n(?=##\s+Question)', content)
        q_blocks = [b for b in blocks if b.strip().startswith("## Question")]
        
        parsed_qs = []
        for block in q_blocks:
            lines = block.strip().split('\n')
            header_line = lines[0]
            
            year_match = re.search(r'\(Year:\s*(\d{4})', header_line)
            year = year_match.group(1) if year_match else None
            
            qid_match = re.search(r'Question ID:\s*([a-zA-Z0-9_-]+)', block, re.IGNORECASE)
            qid = qid_match.group(1) if qid_match else "Unknown"
            
            bold_matches = re.findall(r'\*\*([^*]+)\*\*', block)
            questions = []
            for m in bold_matches:
                m_clean = m.strip()
                if m_clean.startswith("Question ID:") or m_clean.lower().startswith("answer") or m_clean == "Answer" or m_clean == "Answer:":
                    continue
                if len(m_clean) > 20:
                    questions.append(m_clean)
            
            question_text = questions[0] if questions else ""
            if not question_text:
                for line in lines[1:5]:
                    if line.strip().startswith("**") and line.strip().endswith("**"):
                        question_text = line.strip().replace("**", "")
                        break
            
            parsed_qs.append({
                "qid": qid,
                "year": year,
                "original_text": question_text.strip(),
                "tokens": clean_and_tokenize(question_text)
            })
            
        institute_data[inst_name] = parsed_qs
        print(f"Parsed {len(parsed_qs)} questions for {inst_name}")
        
    return institute_data

def set_cell_background(cell, hex_color):
    """Set the cell background color (hex format, e.g. '2C3E50')"""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set the cell padding (in twips: 20 twips = 1 pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def main():
    print("Parsing coaching institute files...")
    institute_data = parse_institute_files()
    
    clean_syllabus_path = os.path.join(syllabus_dir, "GS2_Syllabus_Questions_Formatted.md")
    if not os.path.exists(clean_syllabus_path):
        print(f"Error: Base syllabus file not found: {clean_syllabus_path}")
        return
        
    # Read questions from formatted file
    print("Reading target syllabus questions...")
    with open(clean_syllabus_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    syllabus_questions = []
    current_subject = "Unknown"
    
    i = 0
    while i < len(lines):
        line = lines[i]
        line_strip = line.strip()
        
        if line_strip.startswith("## Subject:"):
            current_subject = line_strip.replace("## Subject:", "").strip()
            i += 1
        elif re.match(r'^Q\d+\.', line_strip):
            question_statement = line_strip
            
            # Read subsequent metadata lines
            metadata_lines = []
            i += 1
            while i < len(lines):
                next_line = lines[i]
                next_line_strip = next_line.strip()
                if not next_line_strip:
                    i += 1
                    continue
                if next_line_strip.startswith('['):
                    metadata_lines.append(next_line_strip)
                    i += 1
                else:
                    break
                    
            combined_metadata = " ".join(metadata_lines)
            year_match = re.search(r'\[Year:\s*(\d{4})\]', combined_metadata)
            year = year_match.group(1) if year_match else "Unknown"
            
            syllabus_questions.append({
                "statement": question_statement,
                "subject": current_subject,
                "year": year,
                "metadata": combined_metadata
            })
        else:
            i += 1
            
    print(f"Found {len(syllabus_questions)} questions in syllabus hierarchy.")
    
    # Audit each question against institutes
    institutes = ["Civilsdaily", "Drishti IAS", "PWOnlyIAS", "Superkalam", "Unacademy"]
    audit_results = []
    
    missing_counts = {inst: 0 for inst in institutes}
    present_counts = {inst: 0 for inst in institutes}
    year_inst_missing = {}
    
    for sq in syllabus_questions:
        q_text = re.sub(r'^Q\d+\.\s*', '', sq["statement"]).strip()
        target_tokens = clean_and_tokenize(q_text)
        year = sq["year"]
        
        statuses = {}
        for inst_name in institutes:
            matched = False
            if inst_name in institute_data:
                inst_qs = institute_data[inst_name]
                # Filter by same year first
                same_year_qs = [iq for iq in inst_qs if iq['year'] == year]
                
                best_match = None
                best_sim = 0.0
                for iq in same_year_qs:
                    sim = jaccard_similarity(target_tokens, iq['tokens'])
                    if sim > best_sim:
                        best_sim = sim
                        best_match = iq
                        
                # Fallback to search all
                if not (best_match and best_sim > 0.35):
                    best_match_any = None
                    best_sim_any = 0.0
                    for iq in inst_qs:
                        sim = jaccard_similarity(target_tokens, iq['tokens'])
                        if sim > best_sim_any:
                            best_sim_any = sim
                            best_match_any = iq
                    if best_match_any and best_sim_any > 0.35:
                        best_match = best_match_any
                        best_sim = best_sim_any
                
                if best_match:
                    matched = True
            
            if matched:
                statuses[inst_name] = "Present"
                present_counts[inst_name] += 1
            else:
                statuses[inst_name] = "Missing"
                missing_counts[inst_name] += 1
                if year != "Unknown":
                    year_inst_missing.setdefault(year, {}).setdefault(inst_name, 0)
                    year_inst_missing[year][inst_name] += 1
                
        audit_results.append({
            "subject": sq["subject"],
            "year": year,
            "statement": sq["statement"],
            "statuses": statuses
        })
        
    print("\nCompilation Audit Statistics:")
    for inst in institutes:
        print(f"  {inst}: Present = {present_counts[inst]}, Missing = {missing_counts[inst]}")
        
    # Generate Word Document
    doc = docx.Document()
    
    # Document Title
    title = doc.add_paragraph()
    run = title.add_run("UPSC Mains GS2: Coaching Institutes Solved Answers Audit")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50) # Dark Slate Blue
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Description
    p_desc = doc.add_paragraph()
    p_desc.add_run("This report identifies the presence or absence of solved answers from 5 coaching institutes (Civilsdaily, Drishti IAS, PWOnlyIAS, Superkalam, and Unacademy) for all GS2 syllabus questions spanning Polity, Governance, Social Justice, and International Relations.")
    
    # Statistics Section
    doc.add_heading("Overall Summary Table", level=2)
    
    stat_table = doc.add_table(rows=1, cols=3)
    stat_table.style = 'Table Grid'
    
    # Format Header Row
    hdr_cells = stat_table.rows[0].cells
    hdr_cells[0].text = "Coaching Institute"
    hdr_cells[1].text = "Solved Answers Present"
    hdr_cells[2].text = "Answers Missing"
    
    for cell in hdr_cells:
        set_cell_background(cell, "2C3E50")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        # format text to bold white
        for p in cell.paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = "Arial"
                r.font.size = Pt(10)
                
    for inst in institutes:
        row_cells = stat_table.add_row().cells
        row_cells[0].text = inst
        row_cells[1].text = str(present_counts[inst])
        row_cells[2].text = str(missing_counts[inst])
        
        # Center align statistics numbers
        row_cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Zebra striping and padding
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
            if idx == 0:
                cell.paragraphs[0].runs[0].bold = True
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
                    
    doc.add_paragraph("") # Spacing
    
    # Year-wise Missing Section
    doc.add_heading("Year-wise Missing Answers Summary", level=2)
    
    year_table = doc.add_table(rows=1, cols=6)
    year_table.style = 'Table Grid'
    
    hdr_cells = year_table.rows[0].cells
    hdr_cells[0].text = "Year"
    for idx, inst in enumerate(institutes):
        hdr_cells[idx+1].text = inst
        
    for cell in hdr_cells:
        set_cell_background(cell, "2C3E50")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            
    # Add Rows
    row_count = 0
    all_years = sorted(list(set(sq["year"] for sq in syllabus_questions if sq["year"] != "Unknown")))
    for yr in all_years:
        row_count += 1
        row_cells = year_table.add_row().cells
        row_cells[0].text = str(yr)
        row_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        for inst_idx, inst_name in enumerate(institutes):
            count = year_inst_missing.get(yr, {}).get(inst_name, 0)
            row_cells[inst_idx+1].text = str(count)
            p = row_cells[inst_idx+1].paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if count > 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) # Red for missing
                else:
                    r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) # Green for 0
                    
        bg_color = "F9FBFB" if row_count % 2 == 0 else "FFFFFF"
        for idx, cell in enumerate(row_cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            if idx == 0:
                cell.paragraphs[0].runs[0].bold = True
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
                    
    doc.add_paragraph("") # Spacing
    
    # Detailed Table Section
    doc.add_heading("Detailed Question Status Table", level=2)
    
    # Create main table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Set widths: Subject (1.0 in), Year (0.5 in), Question (2.5 in), 5 institutes (0.7 in each)
    col_widths = [Inches(1.0), Inches(0.5), Inches(2.5), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7)]
    
    hdr_cells = table.rows[0].cells
    headers = ["Subject", "Year", "Question Statement", "Civilsdaily", "Drishti", "PWOnly", "Kalam", "Unacad."]
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "2C3E50")
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            
    # Add Data Rows
    row_count = 0
    for res in audit_results:
        row_count += 1
        row_cells = table.add_row().cells
        
        # Subject and Year
        row_cells[0].text = res["subject"].title()
        row_cells[1].text = res["year"]
        
        # Center align subject and year
        row_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Question text
        clean_q = re.sub(r'^Q\d+\.\s*', '', res["statement"]).strip()
        if len(clean_q) > 130:
            clean_q = clean_q[:127] + "..."
        row_cells[2].text = clean_q
        
        # Institute columns
        for inst_idx, inst_name in enumerate(institutes):
            status = res["statuses"][inst_name]
            col_idx = 3 + inst_idx
            row_cells[col_idx].text = status
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Format text color and style based on status
            for r in p.runs:
                r.bold = True
                if status == "Present":
                    r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) # Flat Green
                else:
                    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) # Flat Red
                    
        # Apply padding and formatting to all cells in the row
        bg_color = "F9FBFB" if row_count % 2 == 0 else "FFFFFF" # Subtle zebra striping
        for col_idx, cell in enumerate(row_cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(8.5)
                    
    # Force widths on all columns
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    # Save document
    output_path = os.path.join(solved_dir, "gs2_missing_answers_report.docx")
    try:
        doc.save(output_path)
        print(f"Successfully generated and saved audit report to: {output_path}")
    except PermissionError:
        output_path_alt = os.path.join(solved_dir, "gs2_missing_answers_report_updated.docx")
        doc.save(output_path_alt)
        print(f"Permission Denied on main file (likely open in Word). Saved as: {output_path_alt}")

if __name__ == "__main__":
    main()
