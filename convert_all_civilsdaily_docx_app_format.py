import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Directories
civilsdaily_dir = r"C:\Users\Dr. Yogesh\Desktop\mains\neet and upsc cms\upsc\solved paper\civilsdaily"
images_dir = os.path.join(civilsdaily_dir, "images")

def clean_question_text(q_text):
    # Strip leading/trailing asterisks
    q = q_text.strip("*").strip()
    # Remove year and marks bracket patterns
    q = re.sub(r'\[Year:\s*\d{4}\]', '', q)
    q = re.sub(r'\[Marks:\s*\d+\]', '', q)
    return q.strip()

def get_subject_title(subject_key):
    mapping = {
        "ancient_history_and_art__culture": "Ancient History and Art & Culture",
        "science__technology": "Science & Technology",
        "post_independent_india": "Post-Independent India",
        "international_relations": "International Relations",
        "disaster_management": "Disaster Management",
        "economic_development": "Economic Development",
        "environment_and_ecology": "Environment and Ecology",
        "internal_security": "Internal Security",
        "social_justice": "Social Justice",
        "ethics_theoretical_questions": "Ethics Theoretical Questions",
        "ethics_case_studies": "Ethics Case Studies",
        "indian_society": "Indian Society",
        "modern_history": "Modern History",
        "world_history": "World History",
        "geography": "Geography",
        "polity": "Polity",
        "governance": "Governance",
        "agriculture": "Agriculture",
        "social_justice": "Social Justice"
    }
    return mapping.get(subject_key, subject_key.replace("_", " ").title())

def table_to_paragraphs(table_lines):
    clean_lines = [l.strip() for l in table_lines if '---' not in l]
    if not clean_lines:
        return []
    rows = []
    for line in clean_lines:
        parts = [p.strip() for p in line.split('|')]
        if parts and parts[0] == '':
            parts = parts[1:]
        if parts and parts[-1] == '':
            parts = parts[:-1]
        if parts:
            rows.append(parts)
    if not rows:
        return []
    
    out_paras = []
    for row in rows[1:]:
        if len(row) >= 2:
            key = row[0].replace("<br>", " ").replace("<br/>", " ").strip()
            val = row[1].replace("<br>", " ").replace("<br/>", " ").strip()
            out_paras.append(f"- **{key}**: {val}")
        else:
            val = row[0].replace("<br>", " ").replace("<br/>", " ").strip()
            out_paras.append(f"- {val}")
    return out_paras

def convert_md_to_docx(md_path, docx_path, subject_key):
    subject_title = get_subject_title(subject_key)
    print(f"Converting {os.path.basename(md_path)} to {os.path.basename(docx_path)}...")
    
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    # Split by ---
    blocks = re.split(r'\n\s*---\s*\n', content)
    
    doc = Document()
    
    # Page Setup - set default margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    for block in blocks:
        block = block.strip()
        if not block or block.startswith("# UPSC Mains"):
            continue
            
        lines = block.split("\n")
        q_num = None
        q_year = None
        q_paper = None
        q_text = None
        q_marks = None
        
        ans_started = False
        ans_lines = []
        
        idx = 0
        while idx < len(lines):
            line = lines[idx].strip()
            if not line:
                idx += 1
                continue
                
            m_q = re.match(r'^## Question (\d+)\s*\(Year:\s*(\d{4})\s*\|\s*Paper:\s*([^)]+)\)', line)
            if m_q:
                q_num = m_q.group(1)
                q_year = m_q.group(2)
                q_paper = m_q.group(3).strip()
                
                idx += 1
                while idx < len(lines) and not lines[idx].strip():
                    idx += 1
                if idx < len(lines):
                    raw_q_text = lines[idx].strip()
                    m_marks = re.search(r'\[Marks:\s*(\d+)\]', raw_q_text)
                    if m_marks:
                        q_marks = m_marks.group(1)
                    q_text = clean_question_text(raw_q_text)
                idx += 1
                continue
                
            if line.startswith("### Answer"):
                ans_started = True
                idx += 1
                continue
                
            if ans_started:
                ans_lines.append(lines[idx]) # keep original indentation for list prefixing
            idx += 1
            
        if not q_text:
            continue
            
        # Write Question paragraph: Q{num}. {Question Text}
        doc.add_paragraph(f"Q{q_num}. {q_text}")
        doc.add_paragraph("") # Blank line after question
        
        # Parse and process answer paragraphs
        processed_ans = []
        table_buffer = []
        
        for aline in ans_lines:
            aline_strip = aline.strip()
            if not aline_strip:
                continue
                
            # Table detection
            if aline_strip.startswith("|"):
                table_buffer.append(aline)
                continue
            elif table_buffer:
                for t_p in table_to_paragraphs(table_buffer):
                    processed_ans.append((t_p, False))
                table_buffer = []
                
            # Image detection
            if "<img src=" in aline_strip or "![Image]" in aline_strip:
                m_img = re.search(r'src="images/(.*?)"', aline_strip)
                if m_img:
                    filename = m_img.group(1)
                    processed_ans.append((f"<p align=\"center\"><img src=\"images/{filename}\" alt=\"Diagram\" /></p>", False))
                continue
                
            # Headings
            if aline_strip.startswith("## **") and aline_strip.endswith("**"):
                content = aline_strip[5:-2].strip()
                processed_ans.append((f"**{content}**", True))
                continue
            elif aline_strip.startswith("## "):
                content = aline_strip[3:].strip()
                processed_ans.append((f"**{content}**", True))
                continue
                
            # Inline bold check for subheadings
            if aline_strip.startswith("**") and aline_strip.endswith("**") and len(aline_strip) > 4:
                content = aline_strip[2:-2].strip()
                if aline_strip.count("**") == 2:
                    stripped_content = content.lstrip(" '\"“‘")
                    is_subheading = True
                    if content.endswith(".") or content.endswith("?") or content.endswith("!"):
                        is_subheading = False
                    elif re.match(r'^\d+\.', content) or content.startswith("-") or content.startswith("*"):
                        is_subheading = False
                    elif stripped_content and stripped_content[0].islower():
                        is_subheading = False
                    
                    if is_subheading:
                        processed_ans.append((aline_strip, True))
                        continue
                        
            processed_ans.append((aline, False))
            
        if table_buffer:
            for t_p in table_to_paragraphs(table_buffer):
                processed_ans.append((t_p, False))
            
        # Prepend "Explanation: " to the first text paragraph
        first_para_idx = -1
        for i, (p_text, is_subh) in enumerate(processed_ans):
            p_strip = p_text.strip()
            is_bullet = p_strip.startswith("-") or p_strip.startswith("*") or re.match(r'^[a-zA-Z0-9]\.\s+', p_strip) or re.match(r'^\d+\.\s+', p_strip) or re.match(r'^\([a-zA-Z0-9]\)\s+', p_strip)
            is_img = p_strip.startswith("<p align=") or p_strip.startswith("![Image]")
            if not is_bullet and not is_subh and not is_img:
                first_para_idx = i
                break
                
        if first_para_idx != -1:
            orig_text, orig_subh = processed_ans[first_para_idx]
            processed_ans[first_para_idx] = ("Explanation: " + orig_text.strip(), orig_subh)
            
        # Write answer paragraphs to Document
        for p_text, is_subh in processed_ans:
            # If it is an image tag, add text tag and also insert the physical image if exists
            if p_text.startswith("<p align=\"center\"><img src=\"images/"):
                doc.add_paragraph(p_text)
                m_filename = re.search(r'src="images/(.*?)"', p_text)
                if m_filename:
                    filename = m_filename.group(1)
                    img_path = os.path.join(images_dir, filename)
                    if os.path.exists(img_path):
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            p_img.add_run().add_picture(img_path, width=Inches(5.0))
                        except Exception as e:
                            print(f"  Warning: failed to add physical image {filename}: {e}")
                doc.add_paragraph("") # blank line after image
            else:
                p = doc.add_paragraph()
                run = p.add_run(p_text)
                if is_subh:
                    run.font.color.rgb = RGBColor(27, 85, 131) # Royal Blue
                    run.font.bold = True
                
        # Write metadata block
        doc.add_paragraph("") # blank line before metadata
        doc.add_paragraph(f"[Year: {q_year} | Exam: UPSC CSE Mains]")
        marks_part = f" [Marks: {q_marks}]" if q_marks else ""
        doc.add_paragraph(f"[Subject: {subject_title}] [Paper: {q_paper}]{marks_part}")
        doc.add_paragraph("-----")
        doc.add_paragraph("") # blank line after divider
        
    doc.save(docx_path)
    print(f"Saved {os.path.basename(docx_path)}")

def main():
    if not os.path.exists(civilsdaily_dir):
        print(f"Directory {civilsdaily_dir} does not exist.")
        return
        
    # Find all .md files
    for filename in os.listdir(civilsdaily_dir):
        if not filename.endswith(".md"):
            continue
            
        md_path = os.path.join(civilsdaily_dir, filename)
        subject_key = filename.replace(".md", "")
        docx_path = os.path.join(civilsdaily_dir, f"{subject_key}.docx")
        
        convert_md_to_docx(md_path, docx_path, subject_key)

if __name__ == "__main__":
    main()
